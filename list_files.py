from tkinter import messagebox
import docx
from fpdf import XPos, YPos

from orders import agrupaciones_order, secciones_order, papeles_order, all_data_columns_dictionary
from aux_functions import resource_path
from pdf import PDF

# Create Word Document
def listado_word(app, agrupaciones, secciones, papeles, datos, orden, separar, activos, union, numerar, separar_destacados, ancho_columnas, cabecera, archivo):
    agrupaciones, secciones, papeles, activos, datos = format_data(app, agrupaciones, secciones, papeles, separar_destacados, activos, datos)
    doc = docx.Document()
    doc.add_heading(cabecera.strip(), level=1)
    separacion(app, separar, doc, agrupaciones, secciones, papeles, datos, orden, numerar, union, separar_destacados, activos, ancho_columnas)
    output_path = resource_path(f"doc/{archivo}.docx")
    doc.save(output_path)
    messagebox.showinfo("Word creado", f"Archivo creado en {output_path}")

# Create PDF Document
def listado_pdf(app, agrupaciones, secciones, papeles, datos, orden, separar, activos, union, numerar, separar_destacados, ancho_columnas, cabecera, archivo):
    agrupaciones, secciones, papeles, activos, datos = format_data(app, agrupaciones, secciones, papeles, separar_destacados, activos, datos)
    pdf = PDF(cabecera)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_fill_color(194, 194, 194)
    pdf.ln(1)
    pdf.set_y(30)
    separacion(app, separar, pdf, agrupaciones, secciones, papeles, datos, orden, numerar, union, separar_destacados, activos, ancho_columnas)
    output_path = resource_path(f"pdfs/{archivo}.pdf")
    pdf.output(output_path)
    messagebox.showinfo("PDF creado", f"Archivo creado en {output_path}")

# Format data for the report
def format_data(app, agrupaciones, secciones, papeles, separar_destacados, activos, datos):
    if agrupaciones == []:
        agrupaciones = app.agrupaciones
    if secciones == []:
        secciones = app.secciones
    if papeles == []:
        papeles = app.papeles
    agrupaciones = sorted(agrupaciones, key=lambda x: agrupaciones_order.index(x) if x in agrupaciones_order else len(agrupaciones_order))
    secciones = sorted(secciones, key=lambda x: secciones_order.index(x) if x in secciones_order else len(secciones_order))
    papeles = sorted(papeles, key=lambda x: papeles_order.index(x) if x in papeles_order else len(papeles_order))
    # Tratado de los parámetros recibidos
    if activos == "Solo activos":
        activos = ["true"]
    elif activos == "Solo inactivos":
        activos = ["false"]
    else:
        activos = ["true", "false"]
    if separar_destacados:
        datos.append("atril")
    return agrupaciones, secciones, papeles, activos, datos

# Functions to separate the report based on the selected criteria
def separacion(app, separar, doc, agrupaciones, secciones, papeles, datos, orden, numerar, union, separar_destacados, activos, ancho_columnas):
    if separar == "Secciones":
        separar_generico(app, doc, agrupaciones, secciones, papeles, datos, orden, numerar, union, separar_destacados, activos, ancho_columnas, filtro_tipo="seccion", filtro_valores=secciones)
    elif separar == "Agrupaciones":
        separar_generico(app, doc, agrupaciones, secciones, papeles, datos, orden, numerar, union, separar_destacados, activos, ancho_columnas, filtro_tipo="agrupacion", filtro_valores=agrupaciones)
    elif separar == "Papeles":
        separar_generico(app, doc, agrupaciones, secciones, papeles, datos, orden, numerar, union, separar_destacados, activos, ancho_columnas, filtro_tipo="papel", filtro_valores=papeles)
    else:
        separar_generico(app, doc, agrupaciones, secciones, papeles, datos, orden, numerar, union, separar_destacados, activos, ancho_columnas)

def separar_generico(
    app, file, agrupaciones, secciones, papeles, datos, orden, numerar, union,
    separar_destacados, activos, ancho_columnas, filtro_tipo=None, filtro_valores=None
):
    # If any separation is selected
    if filtro_tipo and filtro_valores:
        for valor in filtro_valores:
            filters = [
                ("activo", "in", activos)
            ]
            if filtro_tipo == "seccion":
                filters.append(("papel", "in", papeles))
                filters.append(("agrupacion", "in", agrupaciones))
                filters.append(("seccion", "eq", valor))
            elif filtro_tipo == "agrupacion":
                filters.append(("papel", "in", papeles))
                filters.append(("seccion", "in", secciones))
                filters.append(("agrupacion", "eq", valor))
            elif filtro_tipo == "papel":
                filters.append(("seccion", "in", secciones))
                filters.append(("agrupacion", "in", agrupaciones))
                filters.append(("papel", "eq", valor))

            # Build the query based on the selected filters
            query = app.supabase.table("all_data").select(", ".join(datos))
            for col, op, val in filters:
                if op == "in":
                    query = query.in_(col, val)
                elif op == "eq":
                    query = query.eq(col, val)
            query = query.order(orden)

            # Get the table for the PDF or Word document
            tabla = obtener_tabla_para_archivo( query, datos, numerar, union, separar_destacados)
            if tabla:
                if isinstance(file, docx.document.Document):
                    file.add_heading(valor, level=2)
                    add_text_block_to_doc(file, tabla)
                else:
                    next_table(file, valor, tabla, ancho_columnas)
    # No separation
    else:
        query = app.supabase.table("all_data").select(", ".join(datos))
        query = query.in_("activo", activos)
        query = query.in_("agrupacion", agrupaciones)
        query = query.in_("seccion", secciones)
        query = query.in_("papel", papeles)
        query = query.order(orden)

        # Get the table for the PDF or Word document
        tabla = obtener_tabla_para_archivo(query, datos, numerar, union)
        if tabla:
            if isinstance(file, docx.document.Document):
                add_text_block_to_doc(file, tabla)
            else:
                insertar_tabla(tabla, file, ancho_columnas)

# Function to obtain the table for PDF or Word document
def obtener_tabla_para_archivo(query, datos, numerar, union, separar_destacados=False):
    response = query.execute()
    data_rows = response.data if hasattr(response, "data") else []
    rows = [tuple(row[dato] for dato in datos) for row in data_rows]
    # If there are no rows, return None
    if len(rows) == 0:
        return None
    # Create the table
    tabla = crear_tabla(datos, rows, numerar, union)
    # If there are highlighted rows, separate them
    if separar_destacados:
        aux_tabla = [tabla[0][:-1]]
        between_rows = []
        for item in tabla[1:]:
            if item[-1] == 1:
                aux_tabla.insert(1, item[:-1])
            elif item[-1] == 200:
                aux_tabla.append(item[:-1])
            else:
                between_rows.append(item[:-1])
        for idx, row in enumerate(between_rows):
            aux_tabla.insert(2 + idx, row)
        tabla = aux_tabla
    return tabla

# Function to create the table for the report
def crear_tabla(columns, rows, numerar, union):
    datos = [all_data_columns_dictionary[item].upper() for item in columns]
    cabeceras = ["Nº", *datos]
    j = 1
    tabla = [cabeceras]
    for row in rows:
        tabla.append([j, *row])
        j += 1
    if not numerar:
        tabla = [fila[1:] for fila in tabla]
    if union == "Apellidos,Nombre":
        tabla = unir_apellidos_nombre(numerar, tabla)
    elif union == "Nombre,Apellidos":
        tabla = unir_nombre_apellidos(numerar, tabla)
    return tabla

# Functions to join names and surnames in the table
# Depending on the union type, it will join the names and surnames in the specified order
def unir_nombre_apellidos(numerar, tabla = None):
    if tabla is None:
        return None
    if numerar:
        for fila in tabla:
            fila[1:3] = [" ".join(fila[2:0:-1])]
    elif not numerar:
        for fila in tabla:
            fila[0:2] = [" ".join(fila[1::-1])] # no row number
    return tabla

def unir_apellidos_nombre(numerar, tabla = None):
    if tabla is None:
        return None
    if numerar:
        for fila in tabla:
            fila[1:3] = [", ".join(fila[1:3])]
    elif not numerar:
        for fila in tabla:
            fila[0:2] = [", ".join(fila[0:2])] # no row number
    return tabla

# Function to add a text block to the Word document
def add_text_block_to_doc(doc, tabla):
    for row in tabla[1:]:
        lines = []
        for data in row:
            lines.append(f"{data if data is not None else ''}")
        doc.add_paragraph('\t'.join(lines))
    doc.add_paragraph()

# Function to add a new table to the PDF document
def next_table(pdf, grupo, tabla, ancho_columnas):
    pdf.set_font("Helvetica", "B", 15)
    grupo_w = pdf.get_string_width(grupo) + 6
    if pdf.get_y() + 25 > pdf.page_break_trigger:
        pdf.add_page()
    pdf.cell(grupo_w, 8, grupo, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    insertar_tabla(tabla, pdf, ancho_columnas)

# Function to insert a table into the PDF document
def insertar_tabla(tabla, pdf, anchos):
    longitudes = [int(item) for item in anchos.split(", ")]
    if len(longitudes) != len(tabla[0]):
        messagebox.showerror("Error", "El número de columnas no coincide con el número de anchos especificados para las mismas.")
        raise ValueError("El número de columnas no coincide con el número de anchos especificados para las mismas.")
    else:
        for ancho in longitudes:
            if ancho < 0:
                messagebox.showerror("Error", "Los anchos de las columnas deben ser mayores o iguales a 0.")
                raise ValueError("Los anchos de las columnas deben ser mayores o iguales a 0.")
        pdf.set_font("Helvetica", "", 12)
        elemento(tabla[0], pdf, longitudes, altura=9, fill=True)
        for miembro in tabla[1:]:
            elemento(miembro, pdf, longitudes)
        pdf.ln()
    return

# Function to create a table cell in the PDF document
def elemento(miembro, pdf, longitudes, altura=7, fill=None):
    pdf.set_font('DejaVu', '', 12)
    for i, dato in enumerate(miembro):
        new_x = XPos.RIGHT
        new_y = YPos.TOP
        if dato == miembro[-1]:
            new_x = XPos.LMARGIN
            new_y = YPos.NEXT
        if miembro[i] == None:
            pdf.cell(longitudes[i], altura, "".format(), border=True, new_x=new_x, new_y=new_y, fill=fill)
        else:
            pdf.cell(longitudes[i], altura, "{}".format(miembro[i]), border=True, new_x=new_x, new_y=new_y, fill=fill)
    return